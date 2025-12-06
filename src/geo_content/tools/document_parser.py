"""
Document parsing tool for GEO Content Platform.

Supports PDF and DOCX document parsing for research input.
Handles both local files and S3 URLs.
"""

import logging
import os
import re
import tempfile
from dataclasses import dataclass
from pathlib import Path
from urllib.parse import urlparse

from agents import function_tool

logger = logging.getLogger(__name__)


def is_valid_text_content(content: str, min_readable_ratio: float = 0.7) -> bool:
    """
    Check if content is valid readable text (not binary/garbled).

    Args:
        content: The text content to validate
        min_readable_ratio: Minimum ratio of readable characters (default 0.7)

    Returns:
        True if content appears to be valid text, False otherwise
    """
    if not content or len(content) < 10:
        return False

    # Count printable/readable characters (letters, digits, common punctuation, whitespace)
    # This regex matches most readable text characters across languages
    readable_pattern = re.compile(r'[\w\s.,;:!?\'\"()\[\]{}\-–—/\\@#$%&*+=<>|\u0080-\uFFFF]', re.UNICODE)
    readable_chars = len(readable_pattern.findall(content))
    total_chars = len(content)

    ratio = readable_chars / total_chars if total_chars > 0 else 0

    # Also check for excessive control characters or null bytes
    control_chars = sum(1 for c in content if ord(c) < 32 and c not in '\n\r\t')
    control_ratio = control_chars / total_chars if total_chars > 0 else 0

    if control_ratio > 0.1:  # More than 10% control characters
        logger.warning(f"Content has high control character ratio: {control_ratio:.2%}")
        return False

    if ratio < min_readable_ratio:
        logger.warning(f"Content has low readable character ratio: {ratio:.2%}")
        return False

    return True


def clean_extracted_text(content: str) -> str:
    """
    Clean extracted text by removing problematic characters.

    Args:
        content: Raw extracted text

    Returns:
        Cleaned text content
    """
    if not content:
        return ""

    # Remove null bytes and other problematic control characters
    content = content.replace('\x00', '')

    # Remove other control characters except newlines and tabs
    content = ''.join(c for c in content if ord(c) >= 32 or c in '\n\r\t')

    # Normalize whitespace
    content = re.sub(r'[ \t]+', ' ', content)
    content = re.sub(r'\n{3,}', '\n\n', content)

    return content.strip()

# Try to import document parsing libraries
try:
    from pypdf import PdfReader

    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    PdfReader = None

try:
    from docx import Document as DocxDocument

    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    DocxDocument = None


def _download_s3_file(s3_url: str) -> str | None:
    """
    Download a file from S3 to a temporary location.

    Args:
        s3_url: S3 URL in format s3://bucket/key

    Returns:
        Local temporary file path, or None if download fails
    """
    try:
        import boto3

        parsed = urlparse(s3_url)
        bucket = parsed.netloc
        key = parsed.path.lstrip("/")

        # Get file extension for temp file
        suffix = Path(key).suffix

        # Create temp file with same extension
        fd, temp_path = tempfile.mkstemp(suffix=suffix)
        os.close(fd)

        s3_client = boto3.client("s3")
        s3_client.download_file(bucket, key, temp_path)

        logger.info(f"Downloaded S3 file to: {temp_path}")
        return temp_path

    except Exception as e:
        logger.error(f"Failed to download S3 file {s3_url}: {e}")
        return None


@dataclass
class ParsedDocument:
    """Represents a parsed document."""

    file_path: str
    file_type: str
    title: str
    content: str
    word_count: int
    page_count: int | None
    metadata: dict

    def to_dict(self) -> dict:
        """Convert to dictionary."""
        return {
            "file_path": self.file_path,
            "file_type": self.file_type,
            "title": self.title,
            "content": self.content,
            "word_count": self.word_count,
            "page_count": self.page_count,
            "metadata": self.metadata,
        }


def parse_pdf(file_path: str) -> ParsedDocument | None:
    """
    Parse a PDF document.

    Args:
        file_path: Path to the PDF file

    Returns:
        ParsedDocument or None if parsing fails
    """
    if not PDF_SUPPORT:
        logger.error("PDF support not available. Install pypdf.")
        return None

    try:
        path = Path(file_path)
        if not path.exists():
            logger.error(f"File not found: {file_path}")
            return None

        reader = PdfReader(file_path)

        # Extract text from all pages
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                # Clean each page's text
                cleaned_text = clean_extracted_text(text)
                if cleaned_text:
                    text_parts.append(cleaned_text)

        content = "\n\n".join(text_parts)

        # Validate that we got readable text content
        if not content or not is_valid_text_content(content):
            logger.error(
                f"PDF text extraction failed or produced unreadable content: {file_path}. "
                "The PDF may be a scanned image or have encoding issues."
            )
            return None

        word_count = len(content.split())

        # Extract metadata
        metadata = {}
        if reader.metadata:
            if reader.metadata.title:
                metadata["title"] = reader.metadata.title
            if reader.metadata.author:
                metadata["author"] = reader.metadata.author
            if reader.metadata.subject:
                metadata["subject"] = reader.metadata.subject
            if reader.metadata.creation_date:
                metadata["creation_date"] = str(reader.metadata.creation_date)

        title = metadata.get("title", path.stem)

        return ParsedDocument(
            file_path=file_path,
            file_type="pdf",
            title=title,
            content=content,
            word_count=word_count,
            page_count=len(reader.pages),
            metadata=metadata,
        )

    except Exception as e:
        logger.error(f"Error parsing PDF {file_path}: {e}")
        return None


def parse_docx(file_path: str) -> ParsedDocument | None:
    """
    Parse a DOCX document.

    Args:
        file_path: Path to the DOCX file

    Returns:
        ParsedDocument or None if parsing fails
    """
    if not DOCX_SUPPORT:
        logger.error("DOCX support not available. Install python-docx.")
        return None

    try:
        path = Path(file_path)
        if not path.exists():
            logger.error(f"File not found: {file_path}")
            return None

        doc = DocxDocument(file_path)

        # Extract text from paragraphs
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))

        content = "\n\n".join(text_parts)

        # Clean the extracted content
        content = clean_extracted_text(content)

        # Validate that we got readable text content
        if not content or not is_valid_text_content(content):
            logger.error(
                f"DOCX text extraction failed or produced unreadable content: {file_path}."
            )
            return None

        word_count = len(content.split())

        # Extract metadata from core properties
        metadata = {}
        core_props = doc.core_properties
        if core_props.title:
            metadata["title"] = core_props.title
        if core_props.author:
            metadata["author"] = core_props.author
        if core_props.subject:
            metadata["subject"] = core_props.subject
        if core_props.created:
            metadata["created"] = str(core_props.created)
        if core_props.modified:
            metadata["modified"] = str(core_props.modified)

        title = metadata.get("title", path.stem)

        return ParsedDocument(
            file_path=file_path,
            file_type="docx",
            title=title,
            content=content,
            word_count=word_count,
            page_count=None,  # DOCX doesn't have fixed pages
            metadata=metadata,
        )

    except Exception as e:
        logger.error(f"Error parsing DOCX {file_path}: {e}")
        return None


def parse_text(file_path: str) -> ParsedDocument | None:
    """
    Parse a plain text document.

    Args:
        file_path: Path to the text file

    Returns:
        ParsedDocument or None if parsing fails
    """
    try:
        path = Path(file_path)
        if not path.exists():
            logger.error(f"File not found: {file_path}")
            return None

        # Try different encodings
        content = None
        for encoding in ["utf-8", "utf-16", "latin-1"]:
            try:
                content = path.read_text(encoding=encoding)
                break
            except UnicodeDecodeError:
                continue

        if content is None:
            logger.error(f"Could not decode file: {file_path}")
            return None

        word_count = len(content.split())

        return ParsedDocument(
            file_path=file_path,
            file_type="txt",
            title=path.stem,
            content=content,
            word_count=word_count,
            page_count=None,
            metadata={"encoding": encoding},
        )

    except Exception as e:
        logger.error(f"Error parsing text file {file_path}: {e}")
        return None


def parse_document(file_path: str) -> ParsedDocument | None:
    """
    Parse a document based on its file extension.

    Supports: PDF, DOCX, TXT, MD
    Handles both local file paths and S3 URLs (s3://bucket/key).

    Args:
        file_path: Path to the document (local path or S3 URL)

    Returns:
        ParsedDocument or None if parsing fails or format not supported
    """
    # Handle S3 URLs by downloading to temp file first
    temp_file = None
    original_path = file_path

    if file_path.startswith("s3://"):
        temp_file = _download_s3_file(file_path)
        if not temp_file:
            logger.error(f"Could not download S3 file: {file_path}")
            return None
        local_path = temp_file
    else:
        local_path = file_path

    try:
        path = Path(local_path)
        suffix = path.suffix.lower()

        if suffix == ".pdf":
            result = parse_pdf(local_path)
        elif suffix in [".docx", ".doc"]:
            result = parse_docx(local_path)
        elif suffix in [".txt", ".md", ".markdown"]:
            result = parse_text(local_path)
        else:
            logger.error(f"Unsupported file format: {suffix}")
            result = None

        # Restore original S3 URL in the result
        if result and temp_file:
            result.file_path = original_path

        return result

    finally:
        # Clean up temp file
        if temp_file and Path(temp_file).exists():
            try:
                Path(temp_file).unlink()
                logger.debug(f"Cleaned up temp file: {temp_file}")
            except Exception as e:
                logger.warning(f"Could not delete temp file {temp_file}: {e}")


def parse_documents(file_paths: list[str]) -> list[ParsedDocument]:
    """
    Parse multiple documents.

    Args:
        file_paths: List of file paths to parse

    Returns:
        List of ParsedDocument objects (excludes failed parses)
    """
    results = []
    for file_path in file_paths:
        parsed = parse_document(file_path)
        if parsed:
            results.append(parsed)
    return results


@function_tool
def document_parser_tool(file_path: str) -> dict:
    """
    Parse a document and extract its text content.

    Supports PDF, DOCX, and plain text files (TXT, MD).

    Args:
        file_path: Path to the document file to parse

    Returns:
        Dictionary with file_path, file_type, title, content, word_count,
        page_count, and metadata
    """
    result = parse_document(file_path)
    if result:
        return result.to_dict()
    return {
        "error": f"Failed to parse document: {file_path}",
        "file_path": file_path,
    }


@function_tool
def multi_document_parser_tool(file_paths: list[str]) -> dict:
    """
    Parse multiple documents and extract their text content.

    Supports PDF, DOCX, and plain text files (TXT, MD).

    Args:
        file_paths: List of file paths to parse

    Returns:
        Dictionary with parsed documents and summary statistics
    """
    results = parse_documents(file_paths)
    return {
        "documents": [doc.to_dict() for doc in results],
        "total_documents": len(results),
        "total_words": sum(doc.word_count for doc in results),
        "failed_count": len(file_paths) - len(results),
    }
